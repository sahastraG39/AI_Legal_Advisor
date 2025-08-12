"""
Document Parser Module for AI Legal Document Explainer

This module handles parsing of various document formats including:
- PDF files (using PyMuPDF and pdfplumber)
- Word documents (using python-docx)
- Excel files (using openpyxl)
- Text files
"""

import os
import logging
from typing import Dict, List, Optional, Union
from pathlib import Path

# PDF processing
import fitz  # PyMuPDF
import pdfplumber
from pdf2image import convert_from_path

# Document processing
from docx import Document
import openpyxl

# Data handling
import pandas as pd

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    """Main document parser class for handling various file formats."""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.xlsx': self._parse_excel,
            '.txt': self._parse_text
        }
    
    def parse_document(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """
        Parse a document and return structured data.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing parsed document data
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        logger.info(f"Parsing document: {file_path}")
        
        try:
            parser_func = self.supported_formats[file_extension]
            result = parser_func(file_path)
            
            # Add metadata
            result['file_path'] = str(file_path)
            result['file_size'] = file_path.stat().st_size
            result['file_extension'] = file_extension
            
            logger.info(f"Successfully parsed: {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, any]:
        """Parse PDF documents using multiple methods for comprehensive extraction."""
        result = {
            'content': '',
            'pages': [],
            'metadata': {},
            'text_blocks': [],
            'images': []
        }
        
        try:
            # Method 1: PyMuPDF for basic text and metadata
            with fitz.open(file_path) as doc:
                result['metadata'] = doc.metadata
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    
                    page_data = {
                        'page_number': page_num + 1,
                        'text': text,
                        'bbox': page.rect,
                        'images': len(page.get_images())
                    }
                    result['pages'].append(page_data)
                    result['content'] += text + '\n'
            
            # Method 2: pdfplumber for better text extraction
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    if page.extract_text():
                        # Extract text blocks with positioning
                        text_blocks = page.extract_text_blocks()
                        if text_blocks:
                            result['text_blocks'].extend(text_blocks)
            
            # Method 3: Convert to images for OCR processing
            try:
                images = convert_from_path(file_path, dpi=200)
                result['images'] = [{'page': i+1, 'image': img} for i, img in enumerate(images)]
            except Exception as e:
                logger.warning(f"Could not convert PDF to images: {e}")
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
        
        return result
    
    def _parse_docx(self, file_path: Path) -> Dict[str, any]:
        """Parse Word documents."""
        result = {
            'content': '',
            'paragraphs': [],
            'tables': [],
            'metadata': {}
        }
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    result['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name,
                        'runs': [run.text for run in para.runs if run.text.strip()]
                    })
                    result['content'] += para.text + '\n'
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                result['tables'].append(table_data)
            
            # Extract metadata
            core_props = doc.core_properties
            result['metadata'] = {
                'title': core_props.title,
                'author': core_props.author,
                'subject': core_props.subject,
                'created': core_props.created,
                'modified': core_props.modified
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
        
        return result
    
    def _parse_excel(self, file_path: Path) -> Dict[str, any]:
        """Parse Excel files."""
        result = {
            'sheets': [],
            'content': '',
            'metadata': {}
        }
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = {
                    'name': sheet_name,
                    'rows': [],
                    'max_row': sheet.max_row,
                    'max_column': sheet.max_column
                }
                
                # Extract data from cells
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        sheet_data['rows'].append(row)
                        result['content'] += ' '.join(str(cell) for cell in row if cell) + '\n'
                
                result['sheets'].append(sheet_data)
            
            # Extract metadata
            result['metadata'] = {
                'sheets_count': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames
            }
            
        except Exception as e:
            logger.error(f"Error parsing Excel {file_path}: {e}")
            raise
        
        return result
    
    def _parse_text(self, file_path: Path) -> Dict[str, any]:
        """Parse plain text files."""
        result = {
            'content': '',
            'lines': [],
            'metadata': {}
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                result['content'] = content
                result['lines'] = content.split('\n')
                result['metadata'] = {
                    'line_count': len(result['lines']),
                    'character_count': len(content)
                }
                
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {e}")
            raise
        
        return result
    
    def get_document_summary(self, parsed_data: Dict[str, any]) -> Dict[str, any]:
        """Generate a summary of the parsed document."""
        summary = {
            'file_info': {
                'path': parsed_data.get('file_path', ''),
                'size': parsed_data.get('file_size', 0),
                'format': parsed_data.get('file_extension', '')
            },
            'content_stats': {
                'total_pages': len(parsed_data.get('pages', [])),
                'total_paragraphs': len(parsed_data.get('paragraphs', [])),
                'total_tables': len(parsed_data.get('tables', [])),
                'content_length': len(parsed_data.get('content', ''))
            },
            'extraction_methods': []
        }
        
        # Determine extraction methods used
        if parsed_data.get('pages'):
            summary['extraction_methods'].append('PDF text extraction')
        if parsed_data.get('text_blocks'):
            summary['extraction_methods'].append('PDF text block analysis')
        if parsed_data.get('images'):
            summary['extraction_methods'].append('PDF image conversion')
        if parsed_data.get('paragraphs'):
            summary['extraction_methods'].append('Word document parsing')
        if parsed_data.get('sheets'):
            summary['extraction_methods'].append('Excel spreadsheet parsing')
        
        return summary


def main():
    """Test function for the document parser."""
    parser = DocumentParser()
    
    # Test with a sample file if available
    test_files = [
        'sample.pdf',
        'sample.docx',
        'sample.xlsx',
        'sample.txt'
    ]
    
    for test_file in test_files:
        if Path(test_file).exists():
            try:
                result = parser.parse_document(test_file)
                summary = parser.get_document_summary(result)
                print(f"\nParsed {test_file}:")
                print(f"Summary: {summary}")
            except Exception as e:
                print(f"Error testing {test_file}: {e}")
        else:
            print(f"Test file {test_file} not found")


if __name__ == "__main__":
    main()
